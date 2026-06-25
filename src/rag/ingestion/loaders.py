"""Corpus loaders — Obsidian Markdown vault first, plus generic text and optional binaries.

The PRIMARY corpus is an Obsidian Markdown vault: a folder of ``.md`` notes with YAML
frontmatter, ``[[wikilinks]]``, ``#tags`` and nested folders. The loader also accepts
generic ``.md``/``.markdown``/``.txt`` and OPTIONALLY ``.pdf``/``.docx``. Optional binary
formats import their parser lazily; if the dependency is missing the file is skipped with
a logged warning — ingestion never hard-crashes on a missing optional dependency.

Determinism guarantees:

* the vault is walked with ``sorted()`` (filesystem ``rglob`` order is not stable);
* ``rel_path`` is POSIX-normalized (``as_posix()``) so ids match across Windows and Linux;
* tags and wikilinks are de-duplicated then ``sorted()`` before being stored.
"""

from __future__ import annotations

import datetime as dt
import logging
import re
from pathlib import Path
from typing import Any

from rag.ingestion.models import Document, make_doc_id

logger = logging.getLogger(__name__)

# Suffixes we attempt to load; anything else is skipped.
_MARKDOWN_SUFFIXES = frozenset({".md", ".markdown"})
_TEXT_SUFFIXES = frozenset({".txt"})
_PDF_SUFFIXES = frozenset({".pdf"})
_DOCX_SUFFIXES = frozenset({".docx"})
_ACCEPTED_SUFFIXES = _MARKDOWN_SUFFIXES | _TEXT_SUFFIXES | _PDF_SUFFIXES | _DOCX_SUFFIXES

# Directory names to skip entirely (Obsidian internals, trashed notes, VCS).
_SKIP_DIRS = frozenset({".obsidian", ".trash", ".git"})

_FRONTMATTER_FENCE = "---"
_INLINE_TAG_RE = re.compile(r"(?<!\w)#([A-Za-z0-9_/-]+)")
_WIKILINK_RE = re.compile(r"!?\[\[([^\]]+)\]\]")


def load_corpus(corpus_dir: Path) -> list[Document]:
    """Load every supported file under ``corpus_dir`` into :class:`Document` objects.

    Walks recursively in **sorted** path order (the determinism guarantee), skips Obsidian
    internals and dot-directories, dispatches by suffix, and drops files that a loader
    declines (``None`` — e.g. a missing optional dependency). Returns documents in the same
    sorted path order.
    """
    paths = sorted(p for p in corpus_dir.rglob("*") if p.is_file())
    documents: list[Document] = []
    for path in paths:
        if _is_skipped(path, corpus_dir):
            continue
        document = _load_path(path, corpus_dir)
        if document is not None:
            documents.append(document)
    return documents


def _is_skipped(path: Path, corpus_dir: Path) -> bool:
    """Return ``True`` if ``path`` lives in a skipped directory or has an unsupported suffix."""
    rel_parts = path.relative_to(corpus_dir).parts[:-1]  # exclude the file name itself
    if any(part in _SKIP_DIRS or part.startswith(".") for part in rel_parts):
        return True
    return path.suffix.lower() not in _ACCEPTED_SUFFIXES


def _load_path(path: Path, corpus_dir: Path) -> Document | None:
    """Dispatch a single file to the loader for its suffix."""
    suffix = path.suffix.lower()
    rel_path = path.relative_to(corpus_dir).as_posix()
    if suffix in _MARKDOWN_SUFFIXES:
        return _load_markdown(path, rel_path)
    if suffix in _TEXT_SUFFIXES:
        return _load_text(path, rel_path)
    if suffix in _PDF_SUFFIXES:
        return _load_pdf(path, rel_path)
    if suffix in _DOCX_SUFFIXES:
        return _load_docx(path, rel_path)
    return None


def _read_text(path: Path) -> str:
    """Read a file as UTF-8, replacing undecodable bytes so loading never crashes."""
    return path.read_text(encoding="utf-8", errors="replace")


def _folder_of(rel_path: str) -> str:
    """Return the parent folder of a POSIX ``rel_path`` (``""`` for top-level files)."""
    return rel_path.rsplit("/", 1)[0] if "/" in rel_path else ""


def _build_document(
    path: Path,
    rel_path: str,
    *,
    title: str,
    text: str,
    source_format: str,
    tags: list[str] | None = None,
    wikilinks: list[str] | None = None,
    frontmatter: dict[str, Any] | None = None,
) -> Document:
    """Assemble a :class:`Document` with the standard metadata shape."""
    return Document(
        doc_id=make_doc_id(rel_path),
        source_path=str(path.resolve()),
        rel_path=rel_path,
        title=title,
        text=text,
        metadata={
            "folder": _folder_of(rel_path),
            "tags": tags if tags is not None else [],
            "wikilinks": wikilinks if wikilinks is not None else [],
            "frontmatter": frontmatter if frontmatter is not None else {},
            "source_format": source_format,
        },
    )


def _load_markdown(path: Path, rel_path: str) -> Document:
    """Load an Obsidian/Markdown note: frontmatter, tags, wikilinks, title, body text."""
    raw = _read_text(path)
    frontmatter, body = _split_frontmatter(raw)

    wikilinks, body = _render_wikilinks(body)
    tags = _extract_tags(frontmatter, body)
    title = _resolve_title(frontmatter, body, path)

    return _build_document(
        path,
        rel_path,
        title=title,
        text=body,
        source_format="markdown",
        tags=tags,
        wikilinks=wikilinks,
        frontmatter=frontmatter,
    )


def _load_text(path: Path, rel_path: str) -> Document:
    """Load a plain ``.txt`` file: body is the file text, title is the filename stem."""
    text = _read_text(path)
    return _build_document(
        path,
        rel_path,
        title=path.stem,
        text=text,
        source_format="text",
    )


def _load_pdf(path: Path, rel_path: str) -> Document | None:
    """Load a ``.pdf`` via PyMuPDF (lazy import). Returns ``None`` if PyMuPDF is missing."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.warning("pymupdf not installed; skipping PDF %s", path)
        return None

    parts: list[str] = []
    with fitz.open(path) as pdf:
        for page in pdf:
            parts.append(page.get_text())
    text = "\n\n".join(parts)
    return _build_document(
        path,
        rel_path,
        title=path.stem,
        text=text,
        source_format="pdf",
    )


def _load_docx(path: Path, rel_path: str) -> Document | None:
    """Load a ``.docx`` via python-docx (lazy import). Returns ``None`` if it is missing."""
    try:
        import docx  # python-docx
    except ImportError:
        logger.warning("python-docx not installed; skipping DOCX %s", path)
        return None

    document = docx.Document(str(path))
    text = "\n\n".join(paragraph.text for paragraph in document.paragraphs)
    return _build_document(
        path,
        rel_path,
        title=path.stem,
        text=text,
        source_format="docx",
    )


def _split_frontmatter(raw: str) -> tuple[dict[str, Any], str]:
    """Split YAML frontmatter from the body, returning ``(frontmatter, body)``.

    Frontmatter is only recognized when the file's **first** line is exactly ``---`` and a
    later line is exactly ``---`` (a ``---`` horizontal rule mid-file is not frontmatter).
    With no closing fence the whole file is treated as body (the document is not consumed).
    """
    lines = raw.split("\n")
    if not lines or lines[0].strip() != _FRONTMATTER_FENCE:
        return {}, raw

    closing_index: int | None = None
    for i in range(1, len(lines)):
        if lines[i].strip() == _FRONTMATTER_FENCE:
            closing_index = i
            break
    if closing_index is None:  # no closing fence -> not frontmatter, keep whole body
        return {}, raw

    block = "\n".join(lines[1:closing_index])
    body = "\n".join(lines[closing_index + 1 :])
    return _parse_frontmatter(block), body


def _parse_frontmatter(block: str) -> dict[str, Any]:
    """Parse a YAML frontmatter block into JSON-safe values, never raising.

    Tries PyYAML (lazy import); on a missing dependency or any parse error, falls back to a
    minimal ``key: value`` + simple-list parser. Empty or non-mapping content yields ``{}``.

    PyYAML decodes unquoted dates/datetimes (e.g. the ubiquitous Obsidian
    ``created: 2024-01-15``) into :class:`datetime.date` / :class:`datetime.datetime`
    objects, which are **not** JSON-serializable and would crash ``json.dumps(doc.metadata)``
    in the meta.json / corpus-SHA / Qdrant-payload paths. We therefore coerce the parsed
    result to JSON-safe types (see :func:`_json_safe`) before returning, keeping
    ``Document.metadata`` serializable as the models document.
    """
    if not block.strip():
        return {}
    try:
        import yaml

        parsed = yaml.safe_load(block)
        if isinstance(parsed, dict):
            return _json_safe(parsed)
        return {}
    except Exception:  # noqa: BLE001 - any yaml failure falls back; never crash on odd FM
        return _fallback_frontmatter(block)


def _json_safe(value: Any) -> Any:
    """Recursively coerce ``value`` to JSON-serializable types.

    * ``date`` / ``datetime`` / ``time`` -> ISO-8601 string (``isoformat()``);
    * ``dict`` -> dict with stringified keys and recursively coerced values;
    * ``list`` / ``tuple`` / ``set`` -> list of recursively coerced items (``set`` is
      ``sorted`` first for determinism);
    * native JSON scalars (``str`` / ``int`` / ``float`` / ``bool`` / ``None``) -> unchanged;
    * anything else -> ``str(value)``.

    This guarantees ``json.dumps`` succeeds on the resulting frontmatter (and hence on the
    whole ``Document.metadata`` dict).
    """
    if value is None or isinstance(value, (str, int, float, bool)):
        return value
    if isinstance(value, (dt.datetime, dt.date, dt.time)):
        return value.isoformat()
    if isinstance(value, dict):
        return {str(key): _json_safe(item) for key, item in value.items()}
    if isinstance(value, (list, tuple)):
        return [_json_safe(item) for item in value]
    if isinstance(value, set):
        return [_json_safe(item) for item in sorted(value, key=repr)]
    return str(value)


def _fallback_frontmatter(block: str) -> dict[str, Any]:
    """Minimal frontmatter parser: ``key: value``, inline ``[a, b]`` and ``- item`` lists."""
    result: dict[str, Any] = {}
    current_key: str | None = None
    for line in block.split("\n"):
        if not line.strip():
            continue
        if current_key is not None and line.lstrip().startswith("- "):
            result.setdefault(current_key, [])
            value = result[current_key]
            if isinstance(value, list):
                value.append(line.lstrip()[2:].strip())
            continue
        if ":" not in line:
            continue
        key, _, raw_value = line.partition(":")
        key = key.strip()
        value_text = raw_value.strip()
        if not value_text:
            result[key] = []
            current_key = key
            continue
        result[key] = _coerce_scalar_or_list(value_text)
        current_key = None
    return result


def _coerce_scalar_or_list(value_text: str) -> Any:
    """Coerce a frontmatter scalar; ``[a, b]`` becomes a stripped list, else a string."""
    if value_text.startswith("[") and value_text.endswith("]"):
        inner = value_text[1:-1].strip()
        if not inner:
            return []
        return [item.strip().strip("'\"") for item in inner.split(",")]
    return value_text.strip("'\"")


def _extract_tags(frontmatter: dict[str, Any], body: str) -> list[str]:
    """Union frontmatter ``tags`` with inline ``#tags`` from the body.

    Frontmatter tags may be a list or a comma/space-separated string. Inline tags allow
    nested forms (``#area/sub``). De-duplication preserves first-seen order via
    ``dict.fromkeys`` and the result is ``sorted()`` for determinism. Pure-numeric tokens
    (e.g. ``#123``) are excluded — they are usually issue references, not tags.
    """
    collected: list[str] = []
    collected.extend(_normalize_frontmatter_tags(frontmatter.get("tags")))
    collected.extend(_INLINE_TAG_RE.findall(body))

    deduped = list(dict.fromkeys(collected))
    return sorted(tag for tag in deduped if tag and not tag.isdigit())


def _normalize_frontmatter_tags(value: Any) -> list[str]:
    """Normalize a frontmatter ``tags`` value (list, delimited string, or mapping) to a list.

    Supported shapes:

    * ``list`` -> each item stringified, stripped, leading ``#`` removed;
    * ``str``  -> split on commas/whitespace;
    * ``dict`` -> flatten the **keys** (some vaults emit ``tags:`` as a nested mapping); the
      ``str(dict)`` fallback used to leak a single garbage tag like ``"{'a': 'b'}"``.

    Any other shape is ignored (returns ``[]``) rather than stringified into a pseudo-tag.
    """
    if value is None:
        return []
    if isinstance(value, list):
        return [str(item).strip().lstrip("#") for item in value if str(item).strip()]
    if isinstance(value, str):
        return [token.lstrip("#") for token in re.split(r"[,\s]+", value.strip()) if token]
    if isinstance(value, dict):
        return [str(key).strip().lstrip("#") for key in value if str(key).strip()]
    return []


def _render_wikilinks(body: str) -> tuple[list[str], str]:
    """Replace ``[[target|alias]]``/``[[target]]`` with visible text; collect targets.

    The visible text (alias if present, else target) is kept in the body so the rendered
    string is what chunk offsets index into. ``#section``/``^block`` suffixes are stripped
    from recorded targets. Embeds (``![[...]]``) render the same way (leading ``!`` dropped).
    Targets are returned de-duplicated and ``sorted()``.
    """
    targets: list[str] = []

    def _replace(match: re.Match[str]) -> str:
        inner = match.group(1)
        target_part, _, alias = inner.partition("|")
        target = target_part.split("#", 1)[0].split("^", 1)[0].strip()
        if target:
            targets.append(target)
        visible = alias.strip() if alias else target_part.strip()
        return visible

    rendered = _WIKILINK_RE.sub(_replace, body)
    unique_targets = sorted(dict.fromkeys(targets))
    return unique_targets, rendered


def _resolve_title(frontmatter: dict[str, Any], body: str, path: Path) -> str:
    """Title precedence: frontmatter ``title`` -> first H1 in the body -> filename stem."""
    fm_title = frontmatter.get("title")
    if isinstance(fm_title, str) and fm_title.strip():
        return fm_title.strip()
    h1 = _first_h1(body)
    if h1:
        return h1
    return path.stem


def _first_h1(body: str) -> str | None:
    """Return the first **real** ``# H1`` title, skipping headings inside fenced code blocks.

    This mirrors the chunker's fenced-code-aware heading scan so a ``# comment`` inside a
    ```` ``` ```` fence is never mistaken for the document title (the previous naive regex
    would pick it up). Only an ``# `` at column 0 (with <4 leading spaces, per CommonMark)
    outside a fence counts.
    """
    in_fence = False
    for line in body.split("\n"):
        stripped = line.lstrip()
        if stripped.startswith("```") or stripped.startswith("~~~"):
            in_fence = not in_fence
            continue
        if in_fence:
            continue
        indent = len(line) - len(line.lstrip(" "))
        if indent >= 4:  # indented code block, not a heading
            continue
        if stripped.startswith("# ") and stripped[2:].strip():
            return stripped[2:].strip()
    return None
